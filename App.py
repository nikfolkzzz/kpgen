import tkinter as tk
from tkinter import ttk
from tkinter.constants import VERTICAL 
from Detail import Detail
from NewProj import NewProj




import os
from data import *
from MainMenu import MainMenu

from docx import Document
from docx.shared import Cm, Pt ,Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.enum.table import WD_CELL_ALIGN_VERTICAL



class App(tk.Tk):
    def __init__(self):
        super().__init__() 

        self.geometry('900x600')
        self.arr = exp_type['circ']
        self.all_details = []

        self.draw_num_label = tk.Label(self, text='номер чертежа')
        self.draw_num = tk.Entry(self , width= 15 )

        self.draw_num_label.pack(side=tk.TOP)
        self.draw_num.pack(side=tk.TOP)
        




        self.form_canvas = tk.Canvas(self,bg='lightgreen')
        self.form_canvas.pack(side=tk.LEFT, fill= tk.BOTH , expand=1)

        self.form_frame = tk.Frame(self.form_canvas , bg='lightblue')

        self.scrollbar = tk.Scrollbar(self.form_canvas, orient=VERTICAL, command=self.form_canvas.yview)

        self.scrollbar.pack(side=tk.RIGHT, fill= tk.Y)

        self.form_canvas.configure(yscrollcommand=self.scrollbar.set)


        self.form_canvas.bind('<Configure>', self.task_width)



        self.canvas_frame = self.form_canvas.create_window((0,0), window=self.form_frame, anchor='nw',)

        self.put_details()
        self.put_menu()
#   
# 
        self.btn = ttk.Button(self, text='расчет')
        self.btn.bind("<Button-1>", self.calc_all_forms)
        self.btn.pack(side=tk.TOP,)
# 
# 
# 


        self.bind("<Configure>", self.on_frame_configure)

        self.bind_all("<MouseWheel>", self.mouse_scroll)    




    def put_menu(self):
        self.config(menu = MainMenu(self))

    def put_details(self):



        for idx, d in enumerate(self.arr): 
            detail = Detail(self.form_frame,d)
            self.all_details.append(detail)
            pos_num = divmod(idx,2)[1]
            detail.grid(row=idx, column=0,sticky='ew')
               

                
    def quit(self):
        self.destroy()


    def refresh(self,fej_type):
        self.arr = exp_type[fej_type]


        for f_name in self.all_details:
            self.nametowidget(f_name).destroy()
        self.all_details = []
        self.put_details()




    def task_width(self,e): 
        # canvas_width = e.width
        self.form_canvas.itemconfig(self.canvas_frame, width = e.width)

    def mouse_scroll(self, event): 
            if event.delta: 
                self.form_canvas.yview_scroll(int(-1*(event.delta/120)), "units")
            else: 
                if event.num == 5: 
                    move = 1

                else:
                    move = -1 

                self.form_canvas.yview_scroll(move,"units")


        

    def calc_all_forms(self,evt = None):


        table_rows_arr = []
        for d in self.all_details: 
            table_rows_arr.append(d.func())

        document = Document()

        run = document.add_paragraph().add_run()
        font = run.font
        font.size = Pt(8)

        document.add_heading(f'{self.draw_num.get()}', 1)
 

        table = document.add_table(rows = 1 , cols = 5, )
        
        table.alignment   = WD_ALIGN_PARAGRAPH.CENTER
        table.allow_autofit  = True
        header_cels = table.rows[0].cells

        header_cels[0].text = 'поз.'
        header_cels[1].text = 'кол-во'
        header_cels[2].text = 'Описание'
        header_cels[3].text = 'цена за шт.'
        header_cels[4].text = 'полная стоимость'
        total_price  = [ ]




        for pos, (qty, desc, solo_price, all_price) in enumerate( table_rows_arr ,1 ):

            total_price.append(all_price)
            row_cells = table.add_row().cells
            row_cells[0].text = str(pos)
            row_cells[0].width = Cm(1)   
            row_cells[1].text = str(qty)
            row_cells[1].width = Cm(1)     
            
            row_cells[2].text = str(desc)
            row_cells[3].text = str(solo_price)
            row_cells[4].text = str(all_price)
            row_cells[2].width = Cm(6)

            

        document.add_paragraph(f'полная стоимость: {str(sum(total_price))}')

        

        document.save(f'C://Users//Администратор//Desktop//{self.draw_num.get()}.doc')




    def on_frame_configure(self, event=None):
      self.form_canvas.configure(scrollregion=self.form_canvas.bbox("all"))    




    def create_new_proj(self):
        NewProj(self)




if __name__ == "__main__":
    b = App()

    b.mainloop()
